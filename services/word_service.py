import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MESES = ["enero","febrero","marzo","abril","mayo","junio",
         "julio","agosto","septiembre","octubre","noviembre","diciembre"]

PLANTILLAS = {
    "RE_ASIGNACION": {
        "codigo": "IT-A",
        "asunto": "Revisión de equipo tecnológico disponible para asignación.",
        "antecedentes": (
            "Mediante requerimiento institucional, se solicitó la entrega y configuración "
            "de un equipo de cómputo para <b>{NOMBRES}</b>, con C.I. <b>{CEDULA}</b>, "
            "quien desempeña el cargo de <b>{CARGO}</b> en la <b>{DIRECCION}</b>, a fin "
            "de contar con las herramientas tecnológicas necesarias para el cumplimiento "
            "de sus actividades y responsabilidades institucionales."
        ),
        "desarrollo": [
            "En atención al requerimiento realizado, se efectuó la revisión física y "
            "funcional del equipo tecnológico destinado al personal asignado, verificando "
            "el estado adecuado y funcional del equipo tecnológico.",
            "Como parte del proceso de preparación del equipo, se realizaron las "
            "validaciones necesarias del sistema tecnológico requeridas para el desarrollo "
            "de las actividades laborales.",
            "Se verificó la identificación de los bienes tecnológicos mediante la "
            "validación de sus respectivas etiquetas institucionales y códigos de control interno.",
            "Durante las pruebas y verificaciones efectuadas, se constató que los equipos "
            "presentan un funcionamiento estable y adecuado, cumpliendo con las condiciones "
            "técnicas necesarias para su entrega, asignación y uso dentro de las actividades "
            "institucionales del personal asignado.",
            "Una vez concluidas las verificaciones técnicas y configuraciones "
            "correspondientes, se confirmó que el equipo se encuentra en buenas "
            "condiciones para su utilización.",
        ],
        "conclusiones": [
            "Durante las verificaciones efectuadas se comprobó el correcto funcionamiento "
            "de los recursos tecnológicos requeridos para el desempeño de las funciones "
            "institucionales.",
            "Como resultado de las validaciones efectuadas, se realizó la entrega y "
            "configuración de los bienes tecnológicos a favor de <b>{NOMBRES}</b>, en "
            "su calidad de <b>{CARGO}</b> de la <b>{DIRECCION}</b>, quedando apto "
            "para el desempeño de sus funciones institucionales.",
        ],
    },
    "DESCARGO": {
        "codigo": "IT-D",
        "asunto": "Descargo de bien tecnológico.",
        "antecedentes": (
            "Mediante requerimiento institucional, se solicitó la revisión técnica y "
            "documental de un bien tecnológico asociado a <b>{NOMBRES}</b>, con C.I. "
            "<b>{CEDULA}</b>, quien desempeña el cargo de <b>{CARGO}</b> en la "
            "<b>{DIRECCION}</b>, con la finalidad de sustentar el procedimiento de "
            "descargo correspondiente, de acuerdo con las disposiciones administrativas "
            "y de control institucional vigentes."
        ),
        "desarrollo": [
            "En atención al requerimiento realizado, se efectuó el levantamiento de "
            "información técnica y documental del bien objeto de análisis, revisando "
            "los registros institucionales, códigos patrimoniales y demás elementos "
            "de identificación correspondientes.",
            "Durante la inspección se evaluó la condición actual del bien tecnológico, "
            "recopilando los antecedentes necesarios para determinar su situación dentro "
            "del inventario institucional y sustentar el procedimiento administrativo "
            "correspondiente. Adicionalmente, se documentaron las características "
            "relevantes del bien y su estado de conservación, generando la evidencia "
            "técnica requerida para el trámite de descargo.",
        ],
        "conclusiones": [
            "El análisis efectuado permitió determinar la condición actual del bien "
            "tecnológico y recopilar la información necesaria para sustentar el "
            "procedimiento administrativo correspondiente.",
            "Como resultado de las actividades realizadas, se gestionó el proceso de "
            "descargo del bien tecnológico asociado a <b>{NOMBRES}</b>, en su calidad "
            "de <b>{CARGO}</b> de la <b>{DIRECCION}</b>, de conformidad con la normativa "
            "y procedimientos institucionales vigentes.",
        ],
    },
    "CAMBIO_ACTUALIZACION": {
        "codigo": "IT-CA",
        "asunto": "Cambio y actualización de bien tecnológico.",
        "antecedentes": (
            "Mediante requerimiento institucional, se solicitó la ejecución de actividades "
            "de cambio y actualización de un bien tecnológico asignado a <b>{NOMBRES}</b>, "
            "con C.I. <b>{CEDULA}</b>, quien desempeña el cargo de <b>{CARGO}</b> en la "
            "<b>{DIRECCION}</b>, con el fin de optimizar su desempeño y garantizar la "
            "continuidad de las actividades institucionales."
        ),
        "desarrollo": [
            "En atención al requerimiento realizado, se ejecutaron actividades de cambio "
            "y actualización sobre el bien tecnológico asignado, con la finalidad de "
            "mejorar su desempeño y garantizar la continuidad de las labores institucionales.",
            "Como parte del proceso, se realizaron ajustes técnicos, configuraciones y "
            "adecuaciones orientadas a optimizar las condiciones de uso del equipo de "
            "acuerdo con las necesidades identificadas.",
            "Posteriormente, se efectuaron las pruebas operativas necesarias para "
            "garantizar la correcta implementación de las mejoras realizadas y su "
            "disponibilidad para el usuario asignado.",
        ],
        "conclusiones": [
            "Las actividades ejecutadas permitieron optimizar las condiciones de uso y "
            "desempeño del bien tecnológico, asegurando su adecuada utilización dentro "
            "del entorno institucional.",
            "Como resultado de las acciones efectuadas, se realizó la actualización y "
            "configuración del bien tecnológico asignado a <b>{NOMBRES}</b>, en su calidad "
            "de <b>{CARGO}</b> de la <b>{DIRECCION}</b>, quedando disponible para el "
            "cumplimiento de sus funciones institucionales.",
        ],
    },
    "INSTALACION": {
        "codigo": "IT-I",
        "asunto": "Instalación de bien tecnológico.",
        "antecedentes": (
            "Mediante requerimiento institucional, se solicitó la instalación y "
            "configuración de un bien tecnológico para <b>{NOMBRES}</b>, con C.I. "
            "<b>{CEDULA}</b>, quien desempeña el cargo de <b>{CARGO}</b> en la "
            "<b>{DIRECCION}</b>, a fin de contar con los recursos tecnológicos "
            "necesarios para el cumplimiento de sus actividades y responsabilidades "
            "institucionales."
        ),
        "desarrollo": [
            "En atención al requerimiento realizado, se llevó a cabo la instalación "
            "del bien tecnológico solicitado, considerando los parámetros técnicos y "
            "operativos requeridos para su incorporación al entorno institucional.",
            "Como parte del procedimiento, se efectuaron las configuraciones necesarias "
            "para su puesta en marcha, así como la integración de los recursos asociados "
            "para garantizar su disponibilidad y uso adecuado.",
            "Finalmente, se realizó el registro e identificación del bien mediante la "
            "comprobación de los datos patrimoniales y controles institucionales "
            "correspondientes.",
        ],
        "conclusiones": [
            "Una vez concluido el proceso de instalación, el bien tecnológico quedó "
            "integrado al entorno de trabajo institucional y disponible para su "
            "utilización conforme a los requerimientos del área solicitante.",
            "Como resultado de las actividades realizadas, se efectuó la instalación y "
            "configuración del bien tecnológico a favor de <b>{NOMBRES}</b>, en su "
            "calidad de <b>{CARGO}</b> de la <b>{DIRECCION}</b>, quedando apto para el "
            "desempeño de sus funciones institucionales.",
        ],
    },
    "RE_ESTADO": {
        "codigo": "IT-RE",
        "asunto": "Revisión de estado de bien tecnológico.",
        "antecedentes": (
            "Mediante requerimiento institucional, se solicitó la revisión del estado "
            "físico y operativo de un bien tecnológico asignado a <b>{NOMBRES}</b>, con "
            "C.I. <b>{CEDULA}</b>, quien desempeña el cargo de <b>{CARGO}</b> en la "
            "<b>{DIRECCION}</b>, con el propósito de determinar sus condiciones actuales "
            "de uso y conservación dentro de las actividades institucionales."
        ),
        "desarrollo": [
            "En atención al requerimiento realizado, se efectuó una inspección técnica "
            "del bien tecnológico asignado, con el propósito de determinar sus condiciones "
            "físicas, operativas y de conservación.",
            "Durante la revisión se evaluó el estado general del bien, considerando "
            "aspectos relacionados con su utilización, integridad y desempeño dentro "
            "de las actividades institucionales.",
            "Asimismo, se corroboró la información de identificación patrimonial mediante "
            "la revisión de etiquetas institucionales, números de serie y registros de "
            "control correspondientes.",
        ],
        "conclusiones": [
            "La inspección realizada permitió determinar las condiciones físicas y "
            "operativas del bien tecnológico, evidenciando su estado actual para el "
            "desarrollo de las actividades institucionales.",
            "Como resultado de las acciones efectuadas, se registró el estado del bien "
            "tecnológico asignado a <b>{NOMBRES}</b>, en su calidad de <b>{CARGO}</b> "
            "de la <b>{DIRECCION}</b>, para los fines administrativos y de control "
            "institucional correspondientes.",
        ],
    },
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _fmt_fecha(f):
    if isinstance(f, str):
        f = datetime.strptime(f[:10], "%Y-%m-%d")
    return f"{f.day} de {MESES[f.month - 1]} de {f.year}"


def _sustituir(texto, vd):
    for k, v in vd.items():
        texto = texto.replace("{" + k + "}", str(v) if v else "—")
    return texto


def _add_run_html(para, text, size=11, font="Times New Roman"):
    """Add runs to *para* parsing simple <b>…</b> tags."""
    parts = re.split(r"(<b>|</b>)", text)
    bold = False
    for part in parts:
        if part == "<b>":
            bold = True
        elif part == "</b>":
            bold = False
        elif part:
            run = para.add_run(part)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = font





def _table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _row_height(row, twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for el in trPr.findall(qn("w:trHeight")):
        trPr.remove(el)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(twips))
    trPr.append(trH)


def _add_para(doc, text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              font="Times New Roman", space_before=0, space_after=0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
    return para


def _add_photo(doc, ruta_relativa):
    ruta = os.path.join("static", ruta_relativa)
    if not os.path.exists(ruta):
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(ruta, width=Cm(12))
    except Exception:
        pass


# ── Header / Footer ────────────────────────────────────────────────────────────

def _build_header(doc, config=None):
    if config is None:
        config = {}
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False

    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    img_path = config.get("imagen_encabezado", "")
    if img_path and os.path.exists(img_path):
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(img_path, width=Cm(16.5))
        return

    # Fallback: logos + texto institucional
    tbl = header.add_table(1, 3, width=Cm(16.5))
    _table_no_borders(tbl)

    c0 = tbl.cell(0, 0)
    c0.width = Cm(2.8)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_r = os.path.join("static", "img", "logo_republica.png")
    if os.path.exists(logo_r):
        p0.add_run().add_picture(logo_r, width=Cm(2.2))

    c1 = tbl.cell(0, 1)
    c1.width = Cm(11.0)
    for i, (txt, bold, size) in enumerate([
        ("INSTITUTO NACIONAL DE METEOROLOGÍA E HIDROLOGÍA", True, 9),
        ("INAMHI", True, 9),
        ("DIRECCIÓN ADMINISTRATIVA FINANCIERA", True, 8),
        ("UNIDAD DE TECNOLOGÍAS DE LA INFORMACIÓN Y COMUNICACIÓN", True, 8),
    ]):
        p = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(txt)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"

    c2 = tbl.cell(0, 2)
    c2.width = Cm(2.8)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_n = os.path.join("static", "img", "logo_nuevo_ecuador.png")
    if os.path.exists(logo_n):
        p2.add_run().add_picture(logo_n, width=Cm(2.2))

    tbl2 = header.add_table(1, 3, width=Cm(16.5))
    for row in tbl2.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(8.5)
        row.cells[2].width = Cm(4.0)

    tbl2.cell(0, 0).paragraphs[0].text = "CÓDIGO:"
    for run in tbl2.cell(0, 0).paragraphs[0].runs:
        run.font.size = Pt(8)
        run.font.name = "Arial"

    p_it = tbl2.cell(0, 1).paragraphs[0]
    p_it.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_it = p_it.add_run("INFORME TÉCNICO")
    r_it.bold = True
    r_it.font.size = Pt(9)
    r_it.font.name = "Arial"

    tbl2.cell(0, 2).paragraphs[0].text = "PÁGINA:"
    for run in tbl2.cell(0, 2).paragraphs[0].runs:
        run.font.size = Pt(8)
        run.font.name = "Arial"


def _build_footer(doc, config):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    img_path = config.get("imagen_pie_pagina", "")
    if img_path and os.path.exists(img_path):
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(img_path, width=Cm(16.5))
        return

    # Fallback: texto institucional
    direccion = config.get("footer_direccion", "Av. América N34-61 y Av. Colón, Quito - Ecuador")
    telefono  = config.get("footer_telefono",  "(02) 2261-408")
    web       = config.get("footer_web",       "www.inamhi.gob.ec")

    sep = footer.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(2)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "003b73")
    pBdr.append(top)
    pPr.append(pBdr)

    ft = footer.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft.paragraph_format.space_before = Pt(0)
    ft.paragraph_format.space_after  = Pt(0)
    run = ft.add_run(f"{direccion}   |   Tel: {telefono}   |   {web}")
    run.font.size = Pt(8)
    run.font.name = "Arial"


# ── Blocks ─────────────────────────────────────────────────────────────────────

def _add_metadata_block(doc, asunto, fecha_str):
    # Asunto
    p_a = doc.add_paragraph()
    p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_a.paragraph_format.space_before = Pt(6)
    p_a.paragraph_format.space_after  = Pt(4)
    r_al = p_a.add_run("Asunto: ")
    r_al.bold = True
    r_al.font.size = Pt(11)
    r_al.font.name = "Times New Roman"
    r_av = p_a.add_run(asunto)
    r_av.font.size = Pt(11)
    r_av.font.name = "Times New Roman"

    # Fecha
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_f.paragraph_format.space_before = Pt(0)
    p_f.paragraph_format.space_after  = Pt(6)
    r_fl = p_f.add_run("Fecha: ")
    r_fl.bold = True
    r_fl.font.size = Pt(11)
    r_fl.font.name = "Times New Roman"
    r_fv = p_f.add_run(fecha_str)
    r_fv.font.size = Pt(11)
    r_fv.font.name = "Times New Roman"


def _add_section_title(doc, number, title):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(8)
    run = para.add_run(f"{number}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"


def _add_bien_block(doc, bien):
    tipo_str = (bien.get("tipo_equipo") or "—").upper()

    # Tipo heading
    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_tipo.paragraph_format.space_before = Pt(6)
    p_tipo.paragraph_format.space_after  = Pt(8)
    r_tipo = p_tipo.add_run(tipo_str)
    r_tipo.bold = True
    r_tipo.underline = True
    r_tipo.font.color.rgb = RGBColor(0x00, 0x3b, 0x73)
    r_tipo.font.size = Pt(11)
    r_tipo.font.name = "Times New Roman"

    # Todos los campos como texto plano
    fields = [
        ("MARCA",            "marca"),
        ("ESTADO",           "estado_bien"),
        ("MODELO",           "modelo"),
        ("CÓDIGO ESBYE",     "codigo_esbye"),
        ("NÚMERO DE SERIE",  "serie"),
        ("CÓDIGO ANTERIOR",  "codigo_anterior"),
    ]
    for label, field in fields:
        val = bien.get(field) or ""
        if not val:
            continue
        is_last = (label == "CÓDIGO ANTERIOR")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8) if is_last else Pt(1)
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        r_lbl.font.name = "Times New Roman"
        r_val = p.add_run(val.upper() if label in ("MARCA", "ESTADO") else val)
        r_val.font.size = Pt(10)
        r_val.font.name = "Times New Roman"


def _add_elaborado_aprobado(doc, config):
    doc.add_paragraph()

    elab_nombre = config.get("elaborado_nombre", "TÉCNICO DE TICS")
    elab_cargo  = config.get("elaborado_cargo",  "TÉCNICO")
    apro_nombre = config.get("aprobado_nombre",  "DIRECTOR/A DAF")
    apro_cargo  = config.get("aprobado_cargo",   "DIRECTOR/A")

    tbl = doc.add_table(3, 2)
    tbl.style = "Table Grid"

    # Header row (sin fondo)
    for col_idx, label in enumerate(["ELABORADO POR:", "APROBADO POR:"]):
        cell = tbl.cell(0, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Arial"
    _row_height(tbl.rows[0], 400)

    # Signature space row
    for col_idx in range(2):
        cell = tbl.cell(1, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(" ")
    _row_height(tbl.rows[1], 1800)

    # Name / cargo row
    for col_idx, (nombre, cargo) in enumerate([
        (elab_nombre, elab_cargo),
        (apro_nombre, apro_cargo),
    ]):
        cell = tbl.cell(2, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{nombre}\n{cargo}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
    _row_height(tbl.rows[2], 500)

    for row in tbl.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)


# ── Public API ─────────────────────────────────────────────────────────────────

def generar_word_informe_tecnico(
    tipo, fecha, nombres, cedula, cargo, direccion,
    bienes, fotos_generales, fotos_por_bien,
    numero_informe, ruta_salida, config=None,
    fotos_grupo_2=None,
):
    """
    tipo            : 'RE_ASIGNACION' | 'DESCARGO' | 'RE_ESTADO'
    bienes          : [{'tipo_equipo','marca','modelo','codigo_esbye','serie',
                        'codigo_anterior','estado_bien'}, ...]
    fotos_generales : ['uploads/informes_tecnicos/file.jpg', ...]
    fotos_por_bien  : {0: ['ruta1', ...], 1: [...], ...}   — bien index → rutas
    config          : dict from configuracion_sistema table
    """
    if config is None:
        config = {}

    os.makedirs(os.path.dirname(os.path.abspath(ruta_salida)), exist_ok=True)

    plantilla = PLANTILLAS[tipo]
    vd = {"NOMBRES": nombres, "CEDULA": cedula,
          "CARGO": cargo, "DIRECCION": direccion}
    fecha_str = _fmt_fecha(fecha)

    doc = Document()

    # Márgenes
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(3.5)
    sec.bottom_margin = Cm(3.0)

    _build_header(doc, config)
    _build_footer(doc, config)

    # Título principal
    _add_para(doc, "INFORME TÉCNICO",
              bold=True, size=18,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=4, space_after=4)

    # Código del informe
    _add_para(doc, numero_informe,
              bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=12)

    # Asunto y Fecha (sin tabla)
    _add_metadata_block(doc, plantilla["asunto"], fecha_str)

    _add_para(doc, space_after=4)

    # 1. ANTECEDENTES
    _add_section_title(doc, "1", "ANTECEDENTES")
    ant_text = _sustituir(plantilla["antecedentes"], vd)
    ant_para = doc.add_paragraph()
    ant_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ant_para.paragraph_format.space_before = Pt(2)
    ant_para.paragraph_format.space_after  = Pt(4)
    _add_run_html(ant_para, ant_text)

    # 2. DESARROLLO
    _add_para(doc, space_after=6)
    _add_section_title(doc, "2", "DESARROLLO")
    # (tipo) → (índice donde insertar grupo1, índice donde insertar grupo2)
    _FOTO_POS = {
        "RE_ASIGNACION":       (0, 2),
        "DESCARGO":            (0, 1),
        "RE_ESTADO":           (1, 2),
        "CAMBIO_ACTUALIZACION":(0, 1),
        "INSTALACION":         (0, 2),
    }
    if tipo in _FOTO_POS:
        g1_idx, g2_idx = _FOTO_POS[tipo]
        for idx, dev_txt in enumerate(plantilla["desarrollo"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            _add_run_html(p, _sustituir(dev_txt, vd))
            if idx == g1_idx and fotos_generales:
                _add_para(doc, space_after=4)
                for ruta in fotos_generales:
                    _add_photo(doc, ruta)
            elif idx == g2_idx and fotos_grupo_2:
                _add_para(doc, space_after=4)
                for ruta in fotos_grupo_2:
                    _add_photo(doc, ruta)
    else:
        for dev_txt in plantilla["desarrollo"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            _add_run_html(p, _sustituir(dev_txt, vd))
        if fotos_generales:
            _add_para(doc, space_after=4)
            for ruta in fotos_generales:
                _add_photo(doc, ruta)

    # Subtítulo de bienes (aparece en todos los tipos con fotos por posición)
    if tipo in _FOTO_POS and bienes:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_sub.paragraph_format.space_before = Pt(14)
        p_sub.paragraph_format.space_after  = Pt(4)
        r_sub = p_sub.add_run("Características e identificación del bien tecnológico")
        r_sub.bold = True
        r_sub.font.size = Pt(11)
        r_sub.font.name = "Times New Roman"

    # Bloque por bien
    for idx, bien in enumerate(bienes):
        _add_para(doc, space_after=2)
        _add_bien_block(doc, bien)
        for ruta in fotos_por_bien.get(idx, []):
            _add_photo(doc, ruta)

    _add_para(doc, space_after=4)

    # 3. CONCLUSIONES
    _add_section_title(doc, "3", "CONCLUSIONES")
    for conc_txt in plantilla["conclusiones"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        _add_run_html(p, _sustituir(conc_txt, vd))

    # ELABORADO / APROBADO
    _add_elaborado_aprobado(doc, config)

    doc.save(ruta_salida)
    return ruta_salida
